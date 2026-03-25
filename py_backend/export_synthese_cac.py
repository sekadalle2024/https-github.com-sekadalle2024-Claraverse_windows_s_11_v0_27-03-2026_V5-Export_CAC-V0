"""
Export de Synthèse CAC (Commissaire aux Comptes / Expert-Comptable)
Génère des rapports structurés pour la révision des comptes et le contrôle interne comptable
UTILISE LE TEMPLATE WORD COMME BASE
"""

from fastapi import APIRouter, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from datetime import datetime
import logging
import os
from pathlib import Path
import shutil

logger = logging.getLogger(__name__)
router = APIRouter()

# Chemin vers le template
TEMPLATE_PATH = Path(__file__).parent.parent / "Doc export rapport" / "template final de [Export Synthese CAC].doc"


# === MODÈLES PYDANTIC ===

class FrapPointMetadata(BaseModel):
    """Métadonnées d'un point FRAP"""
    etape: Optional[str] = None
    norme: Optional[str] = None
    methode: Optional[str] = None
    reference: Optional[str] = None


class FrapPoint(BaseModel):
    """Point FRAP (Feuille de Révélation et d'Analyse de Problème)"""
    metadata: FrapPointMetadata
    intitule: str
    observation: Optional[str] = ""
    constat: Optional[str] = ""
    risque: Optional[str] = ""
    recommandation: Optional[str] = ""


class RecosRevisionMetadata(BaseModel):
    """Métadonnées d'un point de recommandation révision"""
    etape: Optional[str] = None
    norme: Optional[str] = None
    methode: Optional[str] = None
    reference: Optional[str] = None


class RecosRevisionPoint(BaseModel):
    """Point de recommandation révision des comptes"""
    metadata: RecosRevisionMetadata
    intitule: str
    description: Optional[str] = ""
    observation: Optional[str] = ""
    ajustement: Optional[str] = ""
    regularisation: Optional[str] = ""


class RecosControleInterneMetadata(BaseModel):
    """Métadonnées d'un point de contrôle interne"""
    etape: Optional[str] = None
    norme: Optional[str] = None
    methode: Optional[str] = None
    reference: Optional[str] = None


class RecosControleInternePoint(BaseModel):
    """Point de recommandation contrôle interne comptable"""
    metadata: RecosControleInterneMetadata
    intitule: str
    observation: Optional[str] = ""
    constat: Optional[str] = ""
    risque: Optional[str] = ""
    recommandation: Optional[str] = ""


class SyntheseCAC_Request(BaseModel):
    """Requête pour générer une synthèse CAC"""
    frap_points: List[FrapPoint] = []
    recos_revision_points: List[RecosRevisionPoint] = []
    recos_controle_interne_points: List[RecosControleInternePoint] = []
    date_rapport: Optional[str] = None
    entite: Optional[str] = "Entité auditée"
    exercice: Optional[str] = None


# === FONCTIONS UTILITAIRES ===

def add_formatted_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False, 
                            font_size: int = 11, indent: float = 0, space_after: int = 6):
    """Ajouter un paragraphe formaté"""
    para = doc.add_paragraph()
    
    if indent > 0:
        para.paragraph_format.left_indent = Inches(indent)
    
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = 1.15
    
    # Gérer les retours à la ligne dans le texte
    text_lines = text.replace('\\n', '\n').replace('\\\\n', '\n').split('\n')
    
    for i, line in enumerate(text_lines):
        if i > 0:
            para.add_run('\n')
        run = para.add_run(line.strip())
        run.font.name = 'Calibri'
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
    
    return para


def add_section_with_label(doc: Document, label: str, content: str, indent: float = 0.25):
    """Ajouter une section avec label en gras et contenu"""
    if not content or content.strip() == "":
        return
    
    para = doc.add_paragraph()
    if indent > 0:
        para.paragraph_format.left_indent = Inches(indent)
    
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.15
    
    # Label en gras
    run_label = para.add_run(f"{label}: ")
    run_label.font.name = 'Calibri'
    run_label.font.size = Pt(11)
    run_label.font.bold = True
    
    # Contenu avec gestion des retours à la ligne
    content_lines = content.replace('\\n', '\n').replace('\\\\n', '\n').split('\n')
    for i, line in enumerate(content_lines):
        if i > 0:
            para.add_run('\n')
        run_content = para.add_run(line.strip())
        run_content.font.name = 'Calibri'
        run_content.font.size = Pt(11)


def add_heading_custom(doc: Document, text: str, level: int = 2, font_size: int = 12):
    """Ajouter un titre personnalisé"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    run = heading.runs[0]
    run.font.name = 'Calibri'
    run.font.size = Pt(font_size)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 56, 100)
    
    return heading


def find_paragraph_with_text(doc: Document, search_text: str):
    """Trouver un paragraphe contenant un texte spécifique"""
    for i, para in enumerate(doc.paragraphs):
        if search_text.lower() in para.text.lower():
            return i, para
    return None, None


def insert_content_after_paragraph(doc: Document, para_index: int, content_func):
    """Insérer du contenu après un paragraphe spécifique"""
    # Cette fonction permet d'insérer du contenu à un endroit précis
    # content_func est une fonction qui ajoute le contenu au document
    content_func()


def create_synthese_cac_from_template(request: SyntheseCAC_Request) -> BytesIO:
    """
    Créer le document Word de synthèse CAC à partir du template
    
    Structure:
    1. INTRODUCTION (déjà dans le template)
    2. OBSERVATIONS D'AUDIT (Recos Révision des Comptes) - À REMPLIR
    3. POINTS DE CONTRÔLE INTERNE (FRAP + Recos Contrôle Interne) - À REMPLIR
    """
    
    # Vérifier que le template existe
    if not TEMPLATE_PATH.exists():
        logger.error(f"❌ Template non trouvé: {TEMPLATE_PATH}")
        raise HTTPException(status_code=500, detail=f"Template non trouvé: {TEMPLATE_PATH}")
    
    # Charger le template
    try:
        doc = Document(str(TEMPLATE_PATH))
        logger.info(f"✅ Template chargé: {TEMPLATE_PATH}")
    except Exception as e:
        logger.error(f"❌ Erreur chargement template: {e}")
        raise HTTPException(status_code=500, detail=f"Erreur chargement template: {e}")
    
    # === REMPLIR LA SECTION 2: OBSERVATIONS D'AUDIT ===
    # Trouver le marqueur "2. OBSERVATIONS D'AUDIT"
    obs_index, obs_para = find_paragraph_with_text(doc, "2. OBSERVATIONS D'AUDIT")
    
    if obs_index is not None:
        logger.info(f"✅ Section 'OBSERVATIONS D'AUDIT' trouvée à l'index {obs_index}")
        
        # Ajouter les points de révision
        if len(request.recos_revision_points) > 0:
            for idx, point in enumerate(request.recos_revision_points, 1):
                # Titre du point
                add_heading_custom(doc, f"2.{idx}. {point.intitule}", level=2, font_size=12)
                
                # Référence
                if point.metadata.reference:
                    add_formatted_paragraph(doc, f"Référence: {point.metadata.reference}", 
                                          italic=True, indent=0.25)
                
                # Description (IMPORTANT: ce champ était manquant)
                if point.description:
                    add_section_with_label(doc, "Description", point.description, indent=0.25)
                
                # Observation
                if point.observation:
                    add_section_with_label(doc, "Observation", point.observation, indent=0.25)
                
                # Ajustement/Reclassement
                if point.ajustement:
                    add_section_with_label(doc, "Ajustement/Reclassement proposé", 
                                         point.ajustement, indent=0.25)
                
                # Régularisation comptable
                if point.regularisation:
                    add_section_with_label(doc, "Régularisation comptable", 
                                         point.regularisation, indent=0.25)
                
                # Espace entre les points
                doc.add_paragraph()
        else:
            add_formatted_paragraph(doc, "Aucune observation d'audit identifiée.", 
                                  italic=True, indent=0.25)
    else:
        logger.warning("⚠️ Section 'OBSERVATIONS D'AUDIT' non trouvée dans le template")
    
    # === REMPLIR LA SECTION 3: POINTS DE CONTRÔLE INTERNE ===
    ci_index, ci_para = find_paragraph_with_text(doc, "3. POINTS DE CONTRÔLE INTERNE")
    
    if ci_index is not None:
        logger.info(f"✅ Section 'POINTS DE CONTRÔLE INTERNE' trouvée à l'index {ci_index}")
        
        # Combiner FRAP et Recos Contrôle Interne
        all_ci_points = []
        
        # Ajouter les FRAP
        for frap in request.frap_points:
            all_ci_points.append({
                'type': 'FRAP',
                'metadata': frap.metadata,
                'intitule': frap.intitule,
                'observation': frap.observation,
                'constat': frap.constat,
                'risque': frap.risque,
                'recommandation': frap.recommandation
            })
        
        # Ajouter les Recos Contrôle Interne
        for reco in request.recos_controle_interne_points:
            all_ci_points.append({
                'type': 'Recos CI',
                'metadata': reco.metadata,
                'intitule': reco.intitule,
                'observation': reco.observation,
                'constat': reco.constat,
                'risque': reco.risque,
                'recommandation': reco.recommandation
            })
        
        if len(all_ci_points) > 0:
            for idx, point in enumerate(all_ci_points, 1):
                # Titre du point
                add_heading_custom(doc, f"3.{idx}. {point['intitule']}", level=2, font_size=12)
                
                # Référence
                if point['metadata'].reference:
                    add_formatted_paragraph(doc, f"Référence: {point['metadata'].reference}", 
                                          italic=True, indent=0.25)
                
                # Type
                add_formatted_paragraph(doc, f"Type: {point['type']}", italic=True, indent=0.25)
                
                # Observation
                if point.get('observation'):
                    add_section_with_label(doc, "Observation", point['observation'], indent=0.25)
                
                # Constat
                if point.get('constat'):
                    add_section_with_label(doc, "Constat", point['constat'], indent=0.25)
                
                # Risques
                if point.get('risque'):
                    add_section_with_label(doc, "Risques identifiés", point['risque'], indent=0.25)
                
                # Recommandation
                if point.get('recommandation'):
                    add_section_with_label(doc, "Recommandation", point['recommandation'], indent=0.25)
                
                # Espace entre les points
                doc.add_paragraph()
        else:
            add_formatted_paragraph(doc, "Aucun point de contrôle interne identifié.", 
                                  italic=True, indent=0.25)
    else:
        logger.warning("⚠️ Section 'POINTS DE CONTRÔLE INTERNE' non trouvée dans le template")
    
    # Sauvegarder dans un buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    logger.info("✅ Document généré avec succès à partir du template")
    
    return buffer


# === ENDPOINT API ===

@router.post("/export-synthese-cac")
async def export_synthese_cac(request: SyntheseCAC_Request):
    """
    Exporter une synthèse CAC en document Word structuré
    UTILISE LE TEMPLATE WORD COMME BASE
    
    Collecte:
    - Points FRAP (contrôle interne opérationnel)
    - Points Recos Révision des Comptes (avec TOUS les champs: description, observation, ajustement, régularisation)
    - Points Recos Contrôle Interne Comptable
    
    Génère un rapport structuré au format CAC/Expert-Comptable
    """
    try:
        total_points = (
            len(request.frap_points) + 
            len(request.recos_revision_points) + 
            len(request.recos_controle_interne_points)
        )
        
        logger.info(f"📊 Export Synthèse CAC: {total_points} points au total")
        logger.info(f"   - FRAP: {len(request.frap_points)}")
        logger.info(f"   - Recos Révision: {len(request.recos_revision_points)}")
        logger.info(f"   - Recos CI: {len(request.recos_controle_interne_points)}")
        
        # Créer le document à partir du template
        doc_buffer = create_synthese_cac_from_template(request)
        
        # Générer le nom de fichier
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = f"synthese_cac_{timestamp}.docx"
        
        logger.info(f"✅ Export réussi: {filename}")
        
        # Retourner le fichier
        return Response(
            content=doc_buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"'
            }
        )
    
    except Exception as e:
        logger.error(f"❌ Erreur export synthèse CAC: {e}")
        raise HTTPException(status_code=500, detail=str(e))
